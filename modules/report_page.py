import streamlit as st
import pandas as pd
import numpy as np
import random
import io
import docx
import re
import time
import logging
from google import genai
from sqlalchemy import text
from modules.database import get_engine
from modules.config import PEMETAAN_WILAYAH

logger = logging.getLogger(__name__)
if not logger.handlers:
    logging.basicConfig(level=logging.INFO)

MONTH_MAP = {'Januari':1, 'Februari':2, 'Maret':3, 'April':4, 'Mei':5, 'Juni':6,
             'Juli':7, 'Agustus':8, 'September':9, 'Oktober':10, 'November':11, 'Desember':12}
INV_MONTH_MAP = {v: k for k, v in MONTH_MAP.items()}

# ==============================================================================
# KONFIGURASI HIERARKI BRS & NORMALISASI
# ==============================================================================
HIERARKI_BRS = {
    "Papua Tengah": {
        "Transportasi Udara": {
            "utama": ['Douw Aturure', 'Mozes Kilangin'],
            "lainnya": ['Enarotali', 'Zugapa Bilorai', 'Moanamani', 'Sinak', 'Illaga', 'Beoga', 'Mulia'],
            "label_subtotal": "Sub Total",
            "label_total": "Total",
            "teks_separator": "Bandara Lainnya"
        },
        "Transportasi Laut": {
            "utama": ['Mimika', 'Nabire'],
            "lainnya": [],
            "label_subtotal": "", 
            "label_total": "Total",
            "teks_separator": ""
        }
    },
    "Papua": {
        "Transportasi Laut": {
            "utama": ['Jayapura', 'Biak'],
            "lainnya": ['Sarmi', 'Serui', 'Waren', 'Kasonaweja'],
            "label_subtotal": "Subtotal",
            "label_total": "Total",
            "teks_separator": "Pelabuhan Lainnya"
        },
        "Transportasi Udara": {
            "utama": ['Sentani', 'Frans Kaisiepo'],
            "lainnya": ['Mararena', 'Stevanus Rumbewas', 'Kasonaweja'],
            "label_subtotal": "Subtotal",
            "label_total": "Total",
            "teks_separator": "Bandara Lainnya"
        }
    },
    "Papua Pegunungan": {
        "Transportasi Udara": {
            "utama": ['Wamena', 'Dekai', 'Batom'],
            "lainnya": ['Oksibil', 'Karubaga'],
            "label_subtotal": "Total",
            "label_total": "Total Keseluruhan",
            "teks_separator": "Bandara Lainnya"
        }
    },
    "Papua Selatan": {
        "Transportasi Laut": {
            "utama": ['Merauke'],
            "lainnya": ['Bade', 'Habesilam', 'Agats', 'Atsy'],
            "label_subtotal": "Jumlah",
            "label_total": "Total",
            "teks_separator": "Pelabuhan lainnya"
        },
        "Transportasi Udara": {
            "utama": ['Moppah'],
            "lainnya": [
                'Okaba', 'Tanah Merah', 'Bomakia', 
                'Mindiptanah', 'Kepi', 'Bade', 
                'Ewer', 'Kamur'
            ],
            "label_subtotal": "Jumlah",
            "label_total": "Total",
            "teks_separator": "Bandara lainnya"
        }
    }
}

def normalisasi_entitas(nama, moda):
    if not isinstance(nama, str):
        return str(nama)
    clean_name = nama.strip()
    if moda == "Transportasi Udara":
        if clean_name.lower().startswith("bandara "):
            clean_name = clean_name[8:].strip()
        mapping_khusus = {"Nabire": "Douw Aturure"}
        return mapping_khusus.get(clean_name, clean_name)
    else:
        if clean_name.lower().startswith("pelabuhan "):
            clean_name = clean_name[10:].strip()
        return clean_name

def build_brs_display_table(report_flat, prov, moda):
    df = report_flat.copy()
    if 'TOTAL' in df.index:
        total_row = df.loc[['TOTAL']]
        df = df.drop(index='TOTAL')
    else:
        total_row = pd.DataFrame()
        
    df = df.reset_index()
    row_col = df.columns[0]
    
    df[row_col] = df[row_col].apply(lambda x: normalisasi_entitas(x, moda))
    df = df.groupby(row_col).sum(min_count=1).reset_index() 
    
    raw_cols = [c for c in df.columns if '(%)' not in c and c != row_col]
    potongan = []
    
    if prov in HIERARKI_BRS and moda in HIERARKI_BRS[prov]:
        config = HIERARKI_BRS[prov][moda]
        df['match_key'] = df[row_col].astype(str).str.lower()
        utama_lower = [str(x).lower() for x in config["utama"]]
        lainnya_lower = [str(x).lower() for x in config["lainnya"]]
        
        df_utama = df[df['match_key'].isin(utama_lower)].copy()
        if not df_utama.empty:
            df_utama = df_utama.drop(columns=['match_key'])
            potongan.append(df_utama)
            if len(config["lainnya"]) > 0:
                sub_utama = pd.DataFrame(df_utama[raw_cols].sum()).T
                sub_utama[row_col] = config["label_subtotal"]
                potongan.append(sub_utama)
            
        if len(config["lainnya"]) > 0:
            separator = pd.DataFrame([{row_col: config["teks_separator"]}])
            for c in df.columns: 
                if c != row_col and c != 'match_key': 
                    separator[c] = np.nan
            potongan.append(separator)
            
            df_lain = df[df['match_key'].isin(lainnya_lower)].copy()
            if not df_lain.empty:
                df_lain = df_lain.drop(columns=['match_key'])
                potongan.append(df_lain)
                sub_lain = pd.DataFrame(df_lain[raw_cols].sum()).T
                sub_lain[row_col] = config["label_subtotal"] + " " 
                potongan.append(sub_lain)
                
        if not potongan and not df.empty:
            if 'match_key' in df.columns: df = df.drop(columns=['match_key'])
            potongan.append(df)
    else:
        if 'match_key' in df.columns: df = df.drop(columns=['match_key'])
        potongan.append(df)
        
    if not total_row.empty:
        t_label = HIERARKI_BRS.get(prov, {}).get(moda, {}).get("label_total", "TOTAL")
        total_row = total_row.reset_index()
        total_row.rename(columns={'index': row_col}, inplace=True)
        total_row[row_col] = t_label
        potongan.append(total_row)
        
    res = pd.concat(potongan, ignore_index=True).set_index(row_col) if potongan else df.set_index(row_col)
    
    col_prev, col_curr = raw_cols[0], raw_cols[1]
    col_cum_prev, col_cum_curr = raw_cols[2], raw_cols[3]
    
    prev_vals_res = res[col_prev].values
    curr_vals_res = res[col_curr].values
    with np.errstate(divide='ignore', invalid='ignore'):
        mtm_res = np.where(prev_vals_res == 0, np.nan, ((curr_vals_res - prev_vals_res) / prev_vals_res) * 100)
    res['M-to-M (%)'] = pd.Series(mtm_res, index=res.index).replace([np.inf, -np.inf], np.nan)
    
    cum_prev_vals_res = res[col_cum_prev].values
    cum_curr_vals_res = res[col_cum_curr].values
    with np.errstate(divide='ignore', invalid='ignore'):
        yoy_res = np.where(cum_prev_vals_res == 0, np.nan, ((cum_curr_vals_res - cum_prev_vals_res) / cum_prev_vals_res) * 100)
    res['Y-on-Y (%)'] = pd.Series(yoy_res, index=res.index).replace([np.inf, -np.inf], np.nan)
    
    return res[[col_prev, col_curr, 'M-to-M (%)', col_cum_prev, col_cum_curr, 'Y-on-Y (%)']]

def get_comparison_data(prov, thn, bln, moda):
    table = "transportasi_udara" if moda == "Transportasi Udara" else "transportasi_laut"
    bln_num = MONTH_MAP[bln]
    thn_int = int(thn)
    engine = get_engine()

    where_clause = "WHERE UPPER(nama_provinsi) = :prov"
    q_curr = f"SELECT * FROM {table} {where_clause} AND CAST(tahun AS TEXT) = :thn AND bulan = :bln"
    df_curr = pd.read_sql(text(q_curr), engine, params={"prov": prov.upper(), "thn": str(thn), "bln": bln})

    prev_bln_num = bln_num - 1 if bln_num > 1 else 12
    prev_thn = thn_int if bln_num > 1 else thn_int - 1
    prev_bln_name = INV_MONTH_MAP[prev_bln_num]
    q_prev = f"SELECT * FROM {table} {where_clause} AND CAST(tahun AS TEXT) = :prev_thn AND bulan = :prev_bln"
    df_prev = pd.read_sql(text(q_prev), engine, params={"prov": prov.upper(), "prev_thn": str(prev_thn), "prev_bln": prev_bln_name})

    months_cum = [INV_MONTH_MAP[i] for i in range(1, bln_num + 1)]
    month_params = {f"m{i}": m for i, m in enumerate(months_cum)}
    months_placeholder = "(" + ", ".join(f":{k}" for k in month_params) + ")"
    q_cum_curr = f"SELECT * FROM {table} {where_clause} AND CAST(tahun AS TEXT) = :thn AND bulan IN {months_placeholder}"
    df_cum_curr = pd.read_sql(text(q_cum_curr), engine, params={"prov": prov.upper(), "thn": str(thn), **month_params})

    q_cum_prev = f"SELECT * FROM {table} {where_clause} AND CAST(tahun AS TEXT) = :prev_thn_full AND bulan IN {months_placeholder}"
    df_cum_prev = pd.read_sql(text(q_cum_prev), engine, params={"prov": prov.upper(), "prev_thn_full": str(thn_int - 1), **month_params})

    return df_curr, df_prev, df_cum_curr, df_cum_prev, prev_bln_name, prev_thn

def format_id_number(x, decimals=2):
    if pd.isna(x) or str(x).lower() in ['nan', 'inf', '-inf', 'undefined']:
        return "Undefined"
    try:
        val = float(x)
        if np.isinf(val) or np.isnan(val): return "Undefined"
        s = f"{val:,.{decimals}f}"
    except (ValueError, TypeError):
        return str(x)
    return s.replace(",", "§").replace(".", ",").replace("§", ".")

NARRATIVE_META = {
    'penumpang_datang':      {'subject': 'Jumlah penumpang yang datang', 'satuan': 'orang', 'is_penumpang': True},
    'penumpang_berangkat':   {'subject': 'Jumlah penumpang yang berangkat', 'satuan': 'orang', 'is_penumpang': True},
    'barang_bongkar_kg':     {'subject': 'Volume barang yang dibongkar', 'satuan': 'kg', 'is_penumpang': False},
    'barang_muat_kg':        {'subject': 'Volume barang yang dimuat', 'satuan': 'kg', 'is_penumpang': False},
    'barang_bongkar_ton':    {'subject': 'Volume barang yang dibongkar', 'satuan': 'ton', 'is_penumpang': False},
    'barang_muat_ton':       {'subject': 'Volume barang yang dimuat', 'satuan': 'ton', 'is_penumpang': False},
    'dn_penumpang_turun':    {'subject': 'Jumlah penumpang yang datang', 'satuan': 'orang', 'is_penumpang': True},
    'dn_penumpang_naik':     {'subject': 'Jumlah penumpang yang berangkat', 'satuan': 'orang', 'is_penumpang': True},
    'dn_bongkar_barang_ton': {'subject': 'Volume barang yang dibongkar', 'satuan': 'ton', 'is_penumpang': False},
    'dn_muat_barang_ton':    {'subject': 'Volume barang yang dimuat', 'satuan': 'ton', 'is_penumpang': False},
}

def _arah_dinamis(pct):
    if pd.isna(pct): return "tercatat"
    if pct > 0: return random.choice(["mengalami lonjakan", "naik", "meningkat", "mengalami pertumbuhan"])
    elif pct < 0: return random.choice(["terkoreksi", "turun", "mengalami penurunan", "menyusut"])
    return "stabil"

def ensure_narasi_cache():
    if "narasi_cache" not in st.session_state:
        st.session_state["narasi_cache"] = {}

def get_cache_key(prov, moda, col_target, bln, thn):
    return f"{prov}|{moda}|{col_target}|{bln}|{thn}"

def get_gemini_api_keys():
    keys = []
    def add_value(v):
        if not v: return
        if isinstance(v, str):
            v = v.strip()
            if v: keys.append(v)
        elif isinstance(v, (list, tuple)):
            for x in v: add_value(x)
        else:
            s = str(v).strip()
            if s: keys.append(s)

    try:
        add_value(st.secrets.get("GEMINI_API_KEYS"))
        add_value(st.secrets.get("GEMINI_API_KEY"))
        add_value(st.secrets.get("GOOGLE_API_KEY"))
        add_value(st.secrets.get("API_GEMINI_KEYS"))
        add_value(st.secrets.get("API_GEMINI_KEY"))
        add_value(st.secrets.get("API-GEMINI-KEYS"))
    except Exception as e:
        logger.exception("Gagal membaca Streamlit secrets: %s", e)

    seen = set()
    unique_keys = []
    for k in keys:
        if k not in seen:
            seen.add(k)
            unique_keys.append(k)
    return unique_keys

def parse_two_paragraphs(text):
    if not text or not str(text).strip(): return None, None
    parts = [p.strip() for p in str(text).strip().split("\n\n") if p.strip()]
    if len(parts) >= 2: return parts[0], "\n\n".join(parts[1:])
    if len(parts) == 1: return parts[0], ""
    return None, None

def generate_single_narrative_ai(df_flat, label, prov, moda, bln, thn, prev_bln, prev_thn):
    cache_key = get_cache_key(prov, moda, label, bln, thn)
    ensure_narasi_cache()
    cache = st.session_state["narasi_cache"]
    
    if cache_key in cache:
        return cache[cache_key], "Cache"

    api_keys = get_gemini_api_keys()
    if not api_keys:
        return None, "No API Key"

    candidate_models = [
        "gemini-2.5-flash",
        "gemini-2.6-flash-lite",
        "gemini-3-flash",
        "gemini-3.1-flash-lite",
        "gemini-3.5-flash",
        "gemini-3.5-flash-lite"
    ]

    if "gemini_key_index" not in st.session_state:
        st.session_state["gemini_key_index"] = 0

    num_keys = len(api_keys)
    data_str = df_flat.to_markdown(index=False)
    
    prompt = (
        "Anda adalah Kepala Pusat Statistik / Penasihat Kebijakan Utama yang menyusun ringkasan eksekutif strategis berstandar tinggi bagi Dewan Pimpinan dan Pengambil Kebijakan.\n"
        f"Buatlah narasi Executive Summary tingkat tinggi yang padat dan tajam (tepat 2 paragraf) untuk indikator statistik \"{label}\" pada moda {moda} Wilayah Provinsi {prov} periode komparasi {bln} {thn} terhadap {prev_bln} {prev_thn}.\n\n"
        "Pedoman & Fokus Penulisan:\n"
        "- Paragraf 1: Analisis komprehensif kinerja bulanan (Month-to-Month/MTM), arah tren sektoral, serta kontribusi agregat dari wilayah-wilayah utama dalam hierarki BRS.\n"
        "- Paragraf 2: Analisis mendalam kinerja kumulatif (Year-to-Date / Year-on-Year), pembacaan deviasi pertumbuhan, serta signifikansi fluktuasi antarwilayah dalam kerangka ekonomi regional.\n"
        "- Gunakan diksi birokratik profesional, objektif, analitis, dengan standarisasi format angka Indonesia.\n"
        "- Jangan sertakan pengantar, sapaan, catatan kaki, ataupun penutup. Langsung berikan 2 paragraf teks yang dipisahkan oleh satu baris kosong (\\n\\n).\n\n"
        "Sumber Data Tabel:\n"
        f"{data_str}"
    )

    for attempt in range(num_keys):
        current_idx = (st.session_state["gemini_key_index"] + attempt) % num_keys
        key = api_keys[current_idx]
        
        for model_name in candidate_models:
            try:
                client = genai.Client(api_key=key.strip())
                response = client.models.generate_content(
                    model=model_name,
                    contents=prompt,
                    config=genai.types.GenerateContentConfig(temperature=0.2)
                )
                raw_text = getattr(response, "text", None)
                if raw_text and str(raw_text).strip():
                    text_clean = str(raw_text).strip()
                    cache[cache_key] = text_clean
                    st.session_state["gemini_key_index"] = (current_idx + 1) % num_keys
                    return text_clean, f"Gemini AI ({model_name})"
            except Exception as e:
                logger.warning("Gagal pada Report dengan Key ke-%d menggunakan model %s: %s. Mencoba opsi model/key lain...", current_idx + 1, model_name, e)
                continue
            
    return None, "Failed"


def generate_narrative_fallback(report_flat, col_target, moda, region_label, bln, thn, prev_bln, prev_thn,
                               col_prev, col_curr, col_cum_prev, col_cum_curr, prov=""):
    meta = NARRATIVE_META.get(col_target, {'subject': 'Total volume', 'satuan': '', 'is_penumpang': False})
    angkutan_kecil = 'udara' if moda == 'Transportasi Udara' else 'laut'
    val_decimals = 0 if meta['satuan'] == 'orang' else 2
    fmt = lambda v: format_id_number(v, decimals=val_decimals)
    fmt_pct = lambda v: format_id_number(v, decimals=2)

    subject = meta['subject']
    if meta['is_penumpang']: 
        subject += f" angkutan {angkutan_kecil} domestik"

    total = report_flat.loc['TOTAL']
    total_curr, total_prev, total_mtm = total[col_curr], total[col_prev], total['M-to-M (%)']
    total_cum_curr, total_cum_prev, total_yoy = total[col_cum_curr], total[col_cum_prev], total['Y-on-Y (%)']

    abs_mtm = fmt_pct(abs(total_mtm) if pd.notna(total_mtm) else np.nan)
    abs_yoy = fmt_pct(abs(total_yoy) if pd.notna(total_yoy) else np.nan)

    p1_options = [
        (
            f"Berdasarkan hasil pencatatan data makro sektoral, {subject.lower()} di Provinsi {prov} pada periode {bln} {thn} "
            f"membukukan realisasi agregat sebesar {fmt(total_curr)} {meta['satuan']}. "
            f"Kinerja bulanan tersebut menunjukkan pergerakan yang {_arah_dinamis(total_mtm)} dengan tingkat fluktuasi "
            f"sebesar {abs_mtm} persen apabila dikomparasikan terhadap baseline operasional bulan {prev_bln} {prev_thn} "
            f"yang berada di angka {fmt(total_prev)} {meta['satuan']}."
        ),
        (
            f"Dinamika sektor transportasi mencatat bahwa {subject.lower()} di wilayah Provinsi {prov} "
            f"mencapai volume total {fmt(total_curr)} {meta['satuan']} selama bulan {bln} {thn}. "
            f"Realisasi ini {_arah_dinamis(total_mtm)} sebesar {abs_mtm} persen secara month-to-month (M-to-M) "
            f"dibandingkan kondisi bulan {prev_bln} {prev_thn} yang mencatatkan angka {fmt(total_prev)} {meta['satuan']}."
        ),
        (
            f"Pada periode {bln} {thn}, aggregate volume {subject.lower()} untuk Provinsi {prov} "
            f"berada pada level {fmt(total_curr)} {meta['satuan']}. "
            f"Perkembangan indikator ini mengindikasikan adanya pergerakan yang {_arah_dinamis(total_mtm)} "
            f"dengan deviasi sebesar {abs_mtm} persen dari performa bulan sebelumnya ({prev_bln} {prev_thn})."
        ),
        (
            f"Meninjau kinerja operasional bulanan, volume {subject.lower()} di Provinsi {prov} pada {bln} {thn} "
            f"tercatat sebesar {fmt(total_curr)} {meta['satuan']}. Capaian tersebut memperlihatkan tren yang {_arah_dinamis(total_mtm)} "
            f"sebesar {abs_mtm} persen jika disandingkan dengan posisi bulan {prev_bln} {prev_thn} "
            f"yang sebelumnya membukukan {fmt(total_prev)} {meta['satuan']}."
        ),
        (
            f"Perkembangan arus {subject.lower()} di Provinsi {prov} pada bulan {bln} {thn} "
            f"menunjukkan angka total mencapai {fmt(total_curr)} {meta['satuan']}. "
            f"Secara bulanan, parameter ini {_arah_dinamis(total_mtm)} dengan persentase perubahan di kisaran {abs_mtm} persen "
            f"terhadap catatan volume pada periode pembanding bulan {prev_bln} {prev_thn}."
        )
    ]

    p2_options = [
        (
            f"Secara kumulatif (Year-to-Date hingga {bln} {thn}), akumulasi realisasi {subject.lower()} "
            f"telah menyentuh angka {fmt(total_cum_curr)} {meta['satuan']}. "
            f"Capaian ini mencatatkan tren pertumbuhan yang {_arah_dinamis(total_yoy)} sebesar {abs_yoy} persen "
            f"secara Year-on-Year (Y-on-Y) jika dibandingkan dengan akumulasi periode yang sama pada tahun sebelumnya "
            f"({fmt(total_cum_prev)} {meta['satuan']}), merefleksikan stabilitas aktivitas ekonomi regional."
        ),
        (
            f"Meninjau kinerja tahun berjalan hingga bulan {bln} {thn}, total realisasi kumulatif tercatat sebesar {fmt(total_cum_curr)} "
            f"{meta['satuan']}. Dibandingkan dengan capaian kumulatif Januari–{bln} tahun sebelumnya ({fmt(total_cum_prev)} {meta['satuan']}), "
            f"indikator ini {_arah_dinamis(total_yoy)} di level {abs_yoy} persen secara Y-on-Y, yang menggambarkan daya tahan "
            f"serta dinamika pemulihan konektivitas wilayah."
        ),
        (
            f"Ditinjau dari perspektif kumulatif tahunan (Januari–{bln} {thn}), volume agregat {subject.lower()} "
            f"mencapai {fmt(total_cum_curr)} {meta['satuan']}. Performa ini menunjukkan kurva laju yang {_arah_dinamis(total_yoy)} "
            f"sebesar {abs_yoy} persen Y-on-Y terhadap baseline operasional tahun sebelumnya, memberi gambaran optimisme "
            f"bagi keberlanjutan moda transportasi di Provinsi {prov}."
        ),
        (
            f"Lebih lanjut, analisis secara kumulatif dari Januari hingga {bln} {thn} menunjukkan total volume penyerapan "
            f"sebesar {fmt(total_cum_curr)} {meta['satuan']}. Angka tersebut mencerminkan dinamika yang {_arah_dinamis(total_yoy)} "
            f"sebesar {abs_yoy} persen secara tahunan (Y-on-Y) apabila dikontraskan dengan capaian periode yang sama tahun lalu "
            f"sebesar {fmt(total_cum_prev)} {meta['satuan']}."
        ),
        (
            f"Dalam rentang waktu tahun berjalan (Year-to-Date) sampai dengan {bln} {thn}, akumulasi arus {subject.lower()} "
            f"terakumulasi pada angka {fmt(total_cum_curr)} {meta['satuan']}. Kinerja makro ini {_arah_dinamis(total_yoy)} "
            f"di level {abs_yoy} persen secara Year-on-Year, menandakan adanya penyesuaian struktural dan pola mobilitas baru "
            f"di kawasan regional."
        )
    ]

    para1 = random.choice(p1_options)
    para2 = random.choice(p2_options)
    
    return para1, para2
                            
def create_complete_master_word_report(prov, thn, bln, all_report_data):
    doc = docx.Document()
    doc.add_heading(f"Laporan Komprehensif Perkembangan Transportasi Provinsi {prov} - {bln} {thn}", level=1)
    doc.add_paragraph(f"Dokumen ini memuat seluruh tabel hierarki BRS dan narasi strategis moda Transportasi Udara dan Transportasi Laut.")
    doc.add_paragraph()

    for item in all_report_data:
        table_no = item['table_no']
        label = item['label']
        moda_name = item['moda']
        p1 = item.get('p1', '')
        p2 = item.get('p2', '')
        df_display = item['df_display']

        angkutan = "Angkutan Udara" if moda_name == "Transportasi Udara" else "Angkutan Laut"
        full_title = f"Tabel {table_no} Perkembangan {label} {angkutan} Dalam Negeri Provinsi {prov}, {bln} {thn}"

        doc.add_heading(full_title, level=2)
        if p1:
            clean_p1 = re.sub(r'^\*\(.*?\)\*\n\n', '', p1)
            doc.add_paragraph(clean_p1)
            doc.add_paragraph()
        
        df_to_export = df_display.reset_index()
        total_rows = len(df_to_export) + 2
        total_cols = len(df_to_export.columns)
        
        table = doc.add_table(rows=total_rows, cols=total_cols)
        table.style = 'Table Grid'
        
        hdr_row_0 = table.rows[0]
        hdr_row_1 = table.rows[1]
        
        for j, col_tuple in enumerate(df_to_export.columns):
            if isinstance(col_tuple, tuple):
                lvl_0 = str(col_tuple[0]).strip()
                lvl_1 = str(col_tuple[1]).strip()
                if j == 0:
                    hdr_row_0.cells[j].text = "Wilayah / Entitas"
                    hdr_row_1.cells[j].text = ""
                else:
                    hdr_row_0.cells[j].text = lvl_0
                    hdr_row_1.cells[j].text = lvl_1 if lvl_1 and lvl_1 != lvl_0 else ""
            else:
                hdr_row_0.cells[j].text = str(col_tuple).strip()
                hdr_row_1.cells[j].text = ""

        try:
            hdr_row_0.cells[0].merge(hdr_row_1.cells[0])
            if total_cols >= 4:
                hdr_row_0.cells[1].merge(hdr_row_0.cells[3])
                if total_cols >= 7:
                    hdr_row_0.cells[4].merge(hdr_row_0.cells[6])
        except Exception:
            pass

        seen_row_contents = set()

        for i, row_data in enumerate(df_to_export.values):
            row_cells = table.rows[i + 2].cells
            first_col_val = str(row_data[0]) if not pd.isna(row_data[0]) else ""
            is_separator_row = "lainnya" in first_col_val.lower()
            
            row_signature = tuple(str(v) for v in row_data)
            if not is_separator_row and row_signature in seen_row_contents and ("subtotal" in first_col_val.lower() or "total" in first_col_val.lower()):
                row_data_cleaned = [""] * len(row_data)
            else:
                seen_row_contents.add(row_signature)
                row_data_cleaned = row_data

            for j, val in enumerate(row_data_cleaned):
                if j == 0:
                    row_cells[j].text = str(val) if not pd.isna(val) else ""
                else:
                    if is_separator_row:
                        row_cells[j].text = ""
                    else:
                        row_cells[j].text = format_id_number(val, decimals=2) if val != "" else ""
            
                         
        doc.add_paragraph()
        if p2:
            clean_p2 = re.sub(r'^\*\(.*?\)\*\n\n', '', p2)
            doc.add_paragraph(clean_p2)
        doc.add_paragraph("---")

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

def prepare_table_item(df_curr, df_prev, df_cum_curr, df_cum_prev, col_target, label, row_col, thn, bln, prev_bln, prev_thn, table_no=None, prov=None, moda=None):
    # Konversi dari Kg ke Ton jika Provinsi Papua Tengah dan Moda Transportasi Udara
    divisor = 1.0
    if prov == "Papua Tengah" and moda == "Transportasi Udara" and "kg" in col_target.lower():
        divisor = 1000.0

    curr_grp = df_curr.groupby(row_col)[col_target].sum() / divisor
    prev_grp = df_prev.groupby(row_col)[col_target].sum() / divisor
    cum_curr_grp = df_cum_curr.groupby(row_col)[col_target].sum() / divisor
    cum_prev_grp = df_cum_prev.groupby(row_col)[col_target].sum() / divisor

    report = pd.DataFrame(index=curr_grp.index)
    col_curr, col_prev = f"{bln} {thn}", f"{prev_bln} {prev_thn}"
    col_cum_curr, col_cum_prev = f"Jan-{bln} {thn}", f"Jan-{bln} {int(thn)-1}"

    report[col_prev] = prev_grp.reindex(report.index).fillna(0)
    report[col_curr] = curr_grp.reindex(report.index).fillna(0)
    
    prev_vals = report[col_prev].values
    curr_vals = report[col_curr].values
    with np.errstate(divide='ignore', invalid='ignore'):
        mtm_pct = np.where(prev_vals == 0, np.nan, ((curr_vals - prev_vals) / prev_vals) * 100)
    report['M-to-M (%)'] = pd.Series(mtm_pct, index=report.index).replace([np.inf, -np.inf], np.nan)

    report[col_cum_prev] = cum_prev_grp.reindex(report.index).fillna(0)
    report[col_cum_curr] = cum_curr_grp.reindex(report.index).fillna(0)
    
    cum_prev_vals = report[col_cum_prev].values
    cum_curr_vals = report[col_cum_curr].values
    with np.errstate(divide='ignore', invalid='ignore'):
        yoy_pct = np.where(cum_prev_vals == 0, np.nan, ((cum_curr_vals - cum_prev_vals) / cum_prev_vals) * 100)
    report['Y-on-Y (%)'] = pd.Series(yoy_pct, index=report.index).replace([np.inf, -np.inf], np.nan)

    sum_prev = report[col_prev].sum()
    sum_curr = report[col_curr].sum()
    sum_cum_prev = report[col_cum_prev].sum()
    sum_cum_curr = report[col_cum_curr].sum()

    total_mtm = ((sum_curr - sum_prev) / sum_prev * 100) if sum_prev != 0 else np.nan
    total_yoy = ((sum_cum_curr - sum_cum_prev) / sum_cum_prev * 100) if sum_cum_prev != 0 else np.nan

    total_row = pd.DataFrame([{
        col_prev: sum_prev, col_curr: sum_curr, 'M-to-M (%)': total_mtm,
        col_cum_prev: sum_cum_prev, col_cum_curr: sum_cum_curr, 'Y-on-Y (%)': total_yoy
    }], index=['TOTAL'])[report.columns]

    report_flat = pd.concat([report, total_row])
    report_display_brs = build_brs_display_table(report_flat, prov, moda)

    cum_label = f"Kumulatif {label}"
    report_display = report_display_brs.copy()
    report_display.columns = pd.MultiIndex.from_tuples([
        (label, col_prev), (label, col_curr), (label, 'M-to-M (%)'),
        (cum_label, col_cum_prev), (cum_label, col_cum_curr), (cum_label, 'Y-on-Y (%)')
    ])
    pct_cols = [(label, 'M-to-M (%)'), (cum_label, 'Y-on-Y (%)')]

    def style_brs_hierarchy(styler):
        def highlight_rows(row):
            styles = [''] * len(row)
            val = str(row.name).strip().lower()
            if val in ['subtotal', 'sub total', 'jumlah']:
                styles = ['font-weight: bold; background-color: #f39c12; color: #ffffff;'] * len(row)
            elif val in ['total', 'total keseluruhan']:
                styles = ['font-weight: bold; background-color: #d35400; color: #ffffff;'] * len(row)
            elif 'lainnya' in val:
                styles = ['font-style: italic; background-color: #e8eaed; color: #000000;'] * len(row)
            return styles
            
        styler = styler.format(format_id_number).background_gradient(subset=pct_cols, cmap='RdYlGn')
        styler = styler.apply(highlight_rows, axis=1)
        return styler

    styled_df = style_brs_hierarchy(report_display.style)

    return {
        'moda': moda,
        'table_no': table_no,
        'label': label,
        'col_target': col_target,
        'prov': prov,
        'bln': bln,
        'thn': thn,
        'prev_bln': prev_bln,
        'prev_thn': prev_thn,
        'col_prev': col_prev,
        'col_curr': col_curr,
        'col_cum_prev': col_cum_prev,
        'col_cum_curr': col_cum_curr,
        'report_flat': report_flat,
        'report_display_brs': report_display_brs,
        'df_display': report_display,
        'styled_df': styled_df,
    }

def render_tables_and_narratives(all_collected_data):
    if not all_collected_data:
        return

    current_moda = None
    for item in all_collected_data:
        if current_moda != item['moda']:
            current_moda = item['moda']
            icon = "✈️" if current_moda == "Transportasi Udara" else "🚢"
            st.subheader(f"{icon} Moda: {current_moda}")

        angkutan = "Angkutan Udara" if current_moda == "Transportasi Udara" else "Angkutan Laut"
        judul = f"Tabel {item['table_no']} Perkembangan {item['label']} {angkutan} Dalam Negeri Provinsi {item['prov']}, {item['bln']} {item['thn']}"
        st.markdown(f"**{judul}**")
        st.dataframe(item['styled_df'], width='stretch')

        region_label = "Bandara" if current_moda == "Transportasi Udara" else "Pelabuhan/Kabupaten"

        h1, h2 = st.columns([5, 1])
        with h1:
            st.markdown(f"**📝 Executive Summary — {item['label']}**")
        with h2:
            if st.button("🔄 Regenerasi", key=f"regen_report_{item['table_no']}", width='stretch'):
                ensure_narasi_cache()
                cache_key = get_cache_key(item['prov'], current_moda, item['label'], item['bln'], item['thn'])
                st.session_state["narasi_cache"].pop(cache_key, None)
                st.session_state["narasi_cache"].pop(cache_key + "|fallback", None)
                st.rerun()

        with st.spinner(f"Menyusun Executive Summary untuk {item['label']}..."):
            text_final, source = generate_single_narrative_ai(
                item['report_display_brs'].reset_index(), 
                item['label'], 
                item['prov'], 
                current_moda, 
                item['bln'], 
                item['thn'], 
                item['prev_bln'], 
                item['prev_thn']
            )

        if text_final:
            p1, p2 = parse_two_paragraphs(text_final)
            p1_text = f"*(Executive Summary - Gemini AI [{source}])*\n\n{p1}" if p1 else ""
            p2_text = p2 if p2 else ""
        else:
            ensure_narasi_cache()
            fallback_cache_key = get_cache_key(item['prov'], current_moda, item['label'], item['bln'], item['thn']) + "|fallback"
            cache = st.session_state["narasi_cache"]
            if fallback_cache_key in cache:
                p1_text, p2_text = cache[fallback_cache_key]
            else:
                p1, p2 = generate_narrative_fallback(
                    report_flat=item['report_flat'],
                    col_target=item['col_target'],
                    moda=current_moda,
                    region_label=region_label,
                    bln=item['bln'],
                    thn=item['thn'],
                    prev_bln=item['prev_bln'],
                    prev_thn=item['prev_thn'],
                    col_prev=item['col_prev'],
                    col_curr=item['col_curr'],
                    col_cum_prev=item['col_cum_prev'],
                    col_cum_curr=item['col_cum_curr'],
                    prov=item['prov']
                )
                p1_text = f"*(Executive Summary - Sistem Fallback)*\n\n{p1}"
                p2_text = p2
                cache[fallback_cache_key] = (p1_text, p2_text)

        if p1_text: st.markdown(p1_text)
        if p2_text: st.markdown(p2_text)
        
        item['p1'] = p1_text
        item['p2'] = p2_text

        st.markdown("---")

def show_report_page():
    st.title("📋 Laporan Komparatif Strategis")
    
    c1, c2, c3 = st.columns(3)
    with c1: prov = st.selectbox("Provinsi", list(PEMETAAN_WILAYAH.keys()))
    with c2: thn = st.selectbox("Tahun", ['2024', '2025', '2026'], index=1)
    with c3: bln = st.selectbox("Bulan", list(MONTH_MAP.keys()))

    if st.button("Generate Semua Laporan (Udara & Laut)"):
        all_collected_data = []
        global_table_counter = 1  
        
        moda_udara = "Transportasi Udara"
        df_cu, df_pr, df_cc, df_cp, p_bln, p_thn = get_comparison_data(prov, thn, bln, moda_udara)
        if not df_cu.empty:
            # Penyesuaian Label untuk Papua Tengah
            if prov == "Papua Tengah":
                targets_udara = [
                    ('penumpang_datang', 'Penumpang Datang'), ('penumpang_berangkat', 'Penumpang Berangkat'),
                    ('barang_bongkar_kg', 'Barang Bongkar (Ton)'), ('barang_muat_kg', 'Barang Muat (Ton)')
                ]
            else:
                targets_udara = [
                    ('penumpang_datang', 'Penumpang Datang'), ('penumpang_berangkat', 'Penumpang Berangkat'),
                    ('barang_bongkar_kg', 'Barang Bongkar (Kg)'), ('barang_muat_kg', 'Barang Muat (Kg)')
                ]

            for col, label in targets_udara:
                item = prepare_table_item(df_cu, df_pr, df_cc, df_cp, col, label, 'nama_bandara', thn, bln, p_bln, p_thn, table_no=global_table_counter, prov=prov, moda=moda_udara)
                all_collected_data.append(item)
                global_table_counter += 1

        moda_laut = "Transportasi Laut"
        df_cu_l, df_pr_l, df_cc_l, df_cp_l, p_bln_l, p_thn_l = get_comparison_data(prov, thn, bln, moda_laut)
        if not df_cu_l.empty:
            row_col_laut = 'nama_kabkota' if prov == "Papua Tengah" else 'nama_pelabuhan'
            targets_laut = [
                ('dn_penumpang_turun', 'Penumpang Turun'), ('dn_penumpang_naik', 'Penumpang Naik'),
                ('dn_bongkar_barang_ton', 'Barang Bongkar (Ton)'), ('dn_muat_barang_ton', 'Barang Muat (Ton)')
            ]
            for col, label in targets_laut:
                item = prepare_table_item(df_cu_l, df_pr_l, df_cc_l, df_cp_l, col, label, row_col_laut, thn, bln, p_bln_l, p_thn_l, table_no=global_table_counter, prov=prov, moda=moda_laut)
                all_collected_data.append(item)
                global_table_counter += 1

        if all_collected_data:
            st.session_state['report_all_data'] = all_collected_data
            st.session_state['report_meta'] = {'prov': prov, 'thn': thn, 'bln': bln}

    if st.session_state.get('report_all_data'):
        all_collected_data = st.session_state['report_all_data']
        meta = st.session_state.get('report_meta', {'prov': prov, 'thn': thn, 'bln': bln})

        render_tables_and_narratives(all_collected_data)

        master_word_file = create_complete_master_word_report(
            meta['prov'], meta['thn'], meta['bln'], all_collected_data
        )
        st.success("Semua data laporan dan narasi berhasil digenerate sepenuhnya!")
        st.download_button(
            label="📥 Download Master Dokumen Word (Semua 8 Tabel & Narasi)",
            data=master_word_file,
            file_name=f"Master_Laporan_Transportasi_{meta['prov']}_{meta['bln']}_{meta['thn']}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
